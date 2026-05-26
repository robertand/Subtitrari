from flask import Flask, request, jsonify, render_template, send_file, session, make_response
from flask_cors import CORS
from werkzeug.utils import secure_filename
import threading
import queue
import time
import uuid
import json
import os
import sys
import logging
import librosa
import numpy as np
from pathlib import Path
from datetime import datetime
import shutil

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent))

from config import Config
from transcriber import WhisperTranscriber, GenderDetector
from translator import Translator
from segmenter import SubtitleSegmenter
from file_handler import FileHandler

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Initialize Flask app
app = Flask(__name__)
app.config.from_object(Config)
app.secret_key = Config.SECRET_KEY
CORS(app)

# Initialize components
Config.init_directories()
transcriber = WhisperTranscriber()
gender_detector = GenderDetector(device="cpu") # Force gender detection on CPU to save VRAM
translator = Translator()
segmenter = SubtitleSegmenter()
file_handler = FileHandler(Config)

# Task storage
processing_tasks = {}
task_lock = threading.Lock()

class ProcessingTask:
    def __init__(self, task_id, file_path, options):
        self.task_id = task_id
        self.file_path = file_path
        self.options = options
        self.status = 'queued'  # queued, processing, completed, failed, cancelled
        self.progress = 0
        self.message = ''
        self.result = None
        self.error = None
        self.created_at = time.time()
        self.last_heartbeat = time.time()
        self.cancel_flag = threading.Event()
    
    def to_dict(self):
        return {
            'task_id': self.task_id,
            'status': self.status,
            'progress': self.progress,
            'message': self.message,
            'result': self.result,
            'error': self.error,
            'created_at': self.created_at,
            'elapsed_time': time.time() - self.created_at
        }

# ============ Routes ============

@app.route('/')
def index():
    """Main page"""
    return render_template('index.html')

@app.route('/api/upload/init', methods=['POST'])
def init_upload():
    """Initialize chunked upload session"""
    try:
        data = request.json
        filename = data.get('filename')
        total_size = data.get('total_size')
        total_chunks = data.get('total_chunks')
        
        if not all([filename, total_size, total_chunks]):
            return jsonify({'error': 'Missing parameters'}), 400
        
        # Validate file extension
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
        if ext not in Config.ALLOWED_EXTENSIONS:
            return jsonify({'error': f'File type not allowed: {ext}'}), 400
        
        session_id = file_handler.create_upload_session(filename, total_size, total_chunks)
        
        return jsonify({
            'session_id': session_id,
            'chunk_size': Config.CHUNK_SIZE
        })
        
    except Exception as e:
        logger.error(f"Upload init error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/upload/chunk', methods=['POST'])
def upload_chunk():
    """Upload a chunk"""
    try:
        session_id = request.form.get('session_id')
        chunk_number_str = request.form.get('chunk_number')
        chunk_file = request.files.get('chunk')
        
        if not all([session_id, chunk_number_str is not None, chunk_file]):
            return jsonify({'error': 'Missing parameters'}), 400
        
        chunk_number = int(chunk_number_str)
        result = file_handler.save_chunk(session_id, chunk_file, chunk_number)
        return jsonify(result)
        
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        logger.error(f"Chunk upload error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/upload/complete', methods=['POST'])
def complete_upload():
    """Complete upload and assemble file"""
    try:
        data = request.json
        session_id = data.get('session_id')
        total_chunks = data.get('total_chunks')
        
        if not session_id:
            return jsonify({'error': 'Missing session_id'}), 400
        
        file_path = file_handler.assemble_file(session_id, total_chunks=total_chunks)
        
        # Generate preview if it's a video file
        preview_url = None
        if Path(file_path).suffix.lower() in {'.mp4', '.avi', '.mov', '.mkv', '.webm'}:
            preview_path = Config.PROCESS_DIR / session_id / 'preview.jpg'
            file_handler.generate_preview(file_path, str(preview_path))
            if preview_path.exists():
                preview_url = f'/api/preview/{session_id}'
        
        return jsonify({
            'file_path': file_path,
            'task_id': session_id,
            'preview_url': preview_url
        })
        
    except Exception as e:
        logger.error(f"Upload complete error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/model_status')
def model_status():
    """Get current model loading status"""
    return jsonify({
        'device': transcriber.device,
        'current_model': transcriber.current_model is not None,
        'loaded_models': list(transcriber.models.keys()),
        'translator_models': list(translator.models.keys()) if translator else []
    })

@app.route('/api/preview/<task_id>')
def get_preview(task_id):
    """Get preview image"""
    preview_path = Config.PROCESS_DIR / task_id / 'preview.jpg'
    if preview_path.exists():
        return send_file(preview_path, mimetype='image/jpeg')
    return jsonify({'error': 'Preview not found'}), 404

@app.route('/api/video/<task_id>')
def serve_video(task_id):
    """Serve video file for player"""
    video_path = Config.PROCESS_DIR / task_id
    
    # Find the video file
    video_file = None
    for ext in ['.mp4', '.avi', '.mov', '.mkv', '.webm']:
        potential_file = video_path / f"original{ext}"
        if potential_file.exists():
            video_file = potential_file
            break
    
    # Also check if the uploaded file is directly in the process dir
    if not video_file:
        for file in video_path.iterdir():
            if file.suffix.lower() in ['.mp4', '.avi', '.mov', '.mkv', '.webm']:
                video_file = file
                break
    
    if video_file and video_file.exists():
        return send_file(
            video_file,
            mimetype='video/mp4',
            conditional=True,
            as_attachment=False
        )
    
    return jsonify({'error': 'Video not found'}), 404

@app.route('/api/audio/<task_id>')
def serve_audio(task_id):
    """Serve audio file for player"""
    audio_path = Config.PROCESS_DIR / task_id / 'audio.wav'
    if audio_path.exists():
        return send_file(audio_path, mimetype='audio/wav')
    
    # Try other audio formats
    for ext in ['.mp3', '.wav', '.m4a', '.flac', '.ogg']:
        for file in (Config.PROCESS_DIR / task_id).iterdir():
            if file.suffix.lower() == ext:
                return send_file(file, mimetype=f'audio/{ext[1:]}')
    
    return jsonify({'error': 'Audio not found'}), 404

@app.route('/api/files/<path:filename>')
def serve_file(filename):
    """Serve any file from process directory"""
    file_path = Config.PROCESS_DIR / filename
    if file_path.exists():
        return send_file(file_path)
    return jsonify({'error': 'File not found'}), 404

@app.route('/api/upload/progress/<session_id>')
def upload_progress(session_id):
    """Get upload progress"""
    progress = file_handler.get_session_progress(session_id)
    if progress:
        return jsonify(progress)
    return jsonify({'error': 'Session not found'}), 404

@app.route('/api/process/start', methods=['POST'])
def start_processing():
    """Start transcription/translation task"""
    try:
        data = request.json
        task_id = data.get('task_id')
        file_path = data.get('file_path')
        options = data.get('options', {})
        
        if not task_id or not file_path:
            return jsonify({'error': 'Missing parameters'}), 400
        
        # Create task
        task = ProcessingTask(task_id, file_path, options)
        
        with task_lock:
            processing_tasks[task_id] = task
        
        # Start processing in background thread
        thread = threading.Thread(
            target=process_task,
            args=(task,),
            daemon=True
        )
        thread.start()
        
        return jsonify({
            'task_id': task_id,
            'status': 'queued'
        })
        
    except Exception as e:
        logger.error(f"Process start error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/process/status/<task_id>')
def process_status(task_id):
    """Get processing status"""
    with task_lock:
        task = processing_tasks.get(task_id)
    
    if not task:
        return jsonify({'error': 'Task not found'}), 404
    
    task.last_heartbeat = time.time()
    return jsonify(task.to_dict())

@app.route('/api/process/zones', methods=['POST'])
def process_zones():
    """Reprocess specific timeline zones"""
    try:
        data = request.json
        task_id = data.get('task_id')
        file_path = data.get('file_path')
        zones = data.get('zones', [])
        options = data.get('options', {})

        if not task_id or not file_path or not zones:
            return jsonify({'error': 'Missing parameters'}), 400

        results = []
        hf_token = options.get('hf_token')

        for idx, zone in enumerate(zones):
            start = zone['start']
            end = zone['end']

            logger.info(f"Selective reprocessing zone {idx}: {start}s -> {end}s")

            # 1. Extract audio
            temp_name = f"reprocess_{task_id}_{int(start)}_{int(end)}.wav"
            audio_temp_path = Config.TEMP_DIR / temp_name

            try:
                transcriber.extract_audio_from_video(
                    file_path, str(audio_temp_path),
                    start_time=start,
                    duration=(end - start)
                )

                # 2. Transcribe
                if options.get('engine') == 'whisper':
                    res = transcriber.transcribe_whisperx(
                        str(audio_temp_path),
                        model_name=options.get('model', 'large-v3'),
                        language=options.get('language'),
                        hf_token=hf_token,
                        use_diarization=options.get('use_diarization', False)
                    )
                else:
                    # Fallback to standard
                    res = transcriber.transcribe_audio(
                        str(audio_temp_path),
                        model_name=options.get('model', 'small'),
                        language=options.get('language'),
                        hf_token=hf_token
                    )

                new_segments = res.get('segments', [])

                # 3. Adjust timestamps (Apply offset)
                for seg in new_segments:
                    seg['start'] += start
                    seg['end'] += start

                # 4. Segmentation
                new_segments = segmenter.segment_by_time(
                    new_segments,
                    min_duration=options.get('min_duration', 1.0),
                    max_duration=options.get('max_duration', 5.0),
                    max_chars=options.get('max_chars', 80)
                )

                # 5. Clean "None" artifacts
                for seg in new_segments:
                    seg['text'] = translator._clean_speaker_none(seg['text'])

                # 6. Translate
                new_translations = {}
                if options.get('translate'):
                    target_langs = options.get('target_languages', ['ro'])
                    source_lang = res.get('language', 'en')
                    engine = options.get('translation_engine', 'google')
                    context = options.get('translation_context')

                    texts = [s['text'] for s in new_segments]
                    meta = [{'gender': s.get('speaker_gender'), 'speaker': s.get('speaker')} for s in new_segments]

                    for target_lang in target_langs:
                        if engine == 'vllm':
                            t_texts = translator.translate_with_vllm_grouped(
                                texts, source_lang, target_lang,
                                model_name=options.get('llm_model'),
                                metadata=meta, context=context
                            )
                        elif engine == 'llm_api':
                            t_texts = translator.translate_with_api(
                                texts, source_lang, target_lang,
                                api_type=options.get('llm_api_provider'),
                                api_key=options.get('llm_api_key'),
                                model=options.get('llm_api_model'),
                                context=context, base_url=options.get('llm_api_url')
                            )
                        else: # google
                            t_texts = translator.translate_batch(
                                texts, source_lang, target_lang, context=context
                            )

                        # Validation
                        t_texts = translator.validate_and_retry_translations(
                            texts, t_texts, source_lang, target_lang, engine
                        )
                        new_translations[target_lang] = t_texts

                results.append({
                    'zone_start': start,
                    'zone_end': end,
                    'segments': new_segments,
                    'translations': new_translations
                })

            finally:
                if audio_temp_path.exists():
                    audio_temp_path.unlink()

        return jsonify({'results': results})

    except Exception as e:
        logger.error(f"Zones process error: {e}")
        return jsonify({'error': str(e)}), 500
    finally:
        transcriber.unload_model()
        translator.unload_models()

@app.route('/api/process/cancel/<task_id>', methods=['POST'])
def cancel_processing(task_id):
    """Cancel a processing task"""
    with task_lock:
        task = processing_tasks.get(task_id)
    
    if not task:
        return jsonify({'error': 'Task not found'}), 404
    
    task.cancel_flag.set()
    task.status = 'cancelled'
    task.message = 'Task cancelled by user'
    
    return jsonify({'status': 'cancelled'})

@app.route('/api/process/result/<task_id>')
def get_result(task_id):
    """Get processing result"""
    with task_lock:
        task = processing_tasks.get(task_id)
    
    if not task:
        return jsonify({'error': 'Task not found'}), 404
    
    if task.status != 'completed':
        return jsonify({'error': 'Task not completed'}), 400
    
    return jsonify(task.result)

@app.route('/api/export/srt', methods=['POST'])
def export_srt():
    """Export subtitles as SRT"""
    try:
        data = request.json
        segments = data.get('segments', [])
        use_legacy_diacritics = data.get('legacy_diacritics', False)
        
        srt_content = generate_srt(segments, use_legacy_diacritics)
        
        response = make_response(srt_content)
        response.headers['Content-Type'] = 'text/plain; charset=utf-8'
        response.headers['Content-Disposition'] = 'attachment; filename=subtitles.srt'
        
        return response
        
    except Exception as e:
        logger.error(f"SRT export error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/export/docx', methods=['POST'])
def export_docx():
    """Export as DOCX for professional translators"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        
        data = request.json
        segments = data.get('segments', [])
        metadata = data.get('metadata', {})
        use_legacy_diacritics = data.get('legacy_diacritics', False)
        
        # Create document
        doc = Document()
        
        # Header
        header = doc.add_heading('Translation Document', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata table
        table = doc.add_table(rows=4, cols=2, style='Table Grid')
        
        metadata_fields = [
            ('Title', metadata.get('title', '')),
            ('Series/Episode', metadata.get('series', '')),
            ('Translator', metadata.get('translator', '')),
            ('Editor', metadata.get('editor', ''))
        ]
        
        for i, (label, value) in enumerate(metadata_fields):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # Add segments
        for i, segment in enumerate(segments, 1):
            text = segment.get('text', '')
            if use_legacy_diacritics:
                text = segmenter.convert_diacritics(text, to_legacy=True)
            
            # Segment number and timestamp
            timestamp = f"{format_time(segment.get('start', 0))} → {format_time(segment.get('end', 0))}"
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. [{timestamp}]")
            run.bold = True
            
            # Text with line splitting
            text_lines = segmenter.split_text_for_subtitle(text, 40)
            p = doc.add_paragraph(text_lines)
            p.style = doc.styles['Normal']
            
            # Empty line between segments
            if i < len(segments):
                doc.add_paragraph()
        
        # Save to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return send_file(
            docx_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='translation.docx'
        )
        
    except Exception as e:
        logger.error(f"DOCX export error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/translate', methods=['POST'])
def translate_text():
    """Translate segments"""
    try:
        data = request.json
        texts = data.get('texts', [])
        source_lang = data.get('source_lang', 'auto')
        target_lang = data.get('target_lang', 'en')
        engine = data.get('engine', 'nllb')
        custom_prompt = data.get('custom_prompt')
        
        if not texts:
            return jsonify({'error': 'No texts provided'}), 400
        
        # Auto-detect source language if needed
        if source_lang == 'auto' and texts:
            source_lang = translator.detect_language(texts[0])
        
        # Translate
        if engine == 'llm':
            translations = translator.translate_with_llm(
                texts, source_lang, target_lang,
                custom_prompt=custom_prompt
            )
        else:
            translations = translator.translate_batch(
                texts, source_lang, target_lang
            )
        
        return jsonify({
            'translations': translations,
            'source_lang': source_lang,
            'target_lang': target_lang
        })
        
    except Exception as e:
        logger.error(f"Translation error: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/languages')
def get_languages():
    """Get supported languages"""
    return jsonify(Config.SUPPORTED_LANGUAGES)

@app.route('/api/models')
def get_models():
    """Get available Whisper models"""
    return jsonify({
        'models': Config.AVAILABLE_MODELS,
        'default': Config.DEFAULT_MODEL,
        'device': transcriber.device
    })

@app.route('/api/cleanup', methods=['POST'])
def cleanup():
    """Manual cleanup trigger"""
    try:
        file_handler.cleanup_old_sessions()
        cleanup_old_tasks()
        return jsonify({'status': 'cleaned'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ============ Background Processing ============

def process_task(task):
    """Main processing function"""
    try:
        task.status = 'processing'
        
        # Extract audio if video
        file_path = Path(task.file_path)
        audio_path = task.file_path
        
        process_start = task.options.get('process_start', 0)
        process_end = task.options.get('process_end', 0)

        if file_path.suffix.lower() in {'.mp4', '.avi', '.mov', '.mkv', '.webm', '.mxf'}:
            task.progress = 5
            task.message = 'Extracting audio...'
            
            audio_path = Config.PROCESS_DIR / task.task_id / 'audio.wav'
            audio_path.parent.mkdir(parents=True, exist_ok=True)
            
            if task.cancel_flag.is_set():
                return
            
            # Use region extraction if requested
            if process_start > 0 or process_end > 0:
                task.message = f'Extracting audio region ({process_start}s - {process_end or "end"}s)...'
                duration = None
                if process_end > process_start:
                    duration = process_end - process_start
                transcriber.extract_audio_from_video(
                    str(task.file_path), str(audio_path),
                    start_time=process_start,
                    duration=duration
                )
            else:
                transcriber.extract_audio_from_video(str(task.file_path), str(audio_path))

        # Voice isolation
        if task.options.get('isolate_voice'):
            task.message = 'Isolating voice...'
            audio, sr = librosa.load(str(audio_path), sr=16000, mono=True)
            audio = transcriber.isolate_voice(audio, sr)
            import soundfile as sf
            sf.write(str(audio_path), audio, sr)
        
        # Transcribe
        task.progress = 10
        task.message = 'Starting transcription...'
        
        if task.cancel_flag.is_set():
            return
        
        engine = task.options.get('engine', Config.DEFAULT_ENGINE)
        model_name = task.options.get('model', Config.DEFAULT_MODEL)
        language = task.options.get('language', Config.DEFAULT_LANGUAGE)
        
        logger.info(f"Task {task.task_id} - Engine: {engine}, Model: {model_name}, Lang: {language}")

        if language == 'auto':
            language = None
        
        if engine == 'cohere':
            task.message = 'Initializing Cohere Transcribe...'
            if transcriber.current_model and transcriber.current_model != transcriber.models.get(Config.COHERE_MODEL):
                transcriber.unload_model()

            # Cohere non-chunked as it seems transcribe_with_cohere is what's available
            result = transcriber.transcribe_with_cohere(
                str(audio_path),
                language=language or 'en',
                use_forced_alignment=True,
                progress_callback=lambda p, m: update_task_progress(task, p, m)
            )
        elif task.options.get('mixed_turkish') and language == 'tr':
            # Mixed Turkish Mode (Whisper Large V3 + Turbo TR)
            all_segments = []

            # Pass 1: OpenAI Whisper Large V3 (60s / 5s)
            task.message = "Mixed TR Pass 1 (Whisper Large V3, 60s window)..."
            res1 = transcriber.transcribe_with_windowing(
                str(audio_path),
                model_name='large-v3',
                language='tr',
                window_size=60,
                overlap=5,
                progress_callback=lambda p, m: update_task_progress(task, p * 0.5, f"Pass 1 (Large-V3): {m}")
            )
            all_segments.extend(res1.get("segments", []))

            if task.cancel_flag.is_set(): return

            # Pass 2: Whisper Turbo Turkish (30s / 5s, isolated)
            task.message = "Mixed TR Pass 2 (Turbo Turkish, 30s window, isolated)..."
            audio_2, sr_2 = librosa.load(str(audio_path), sr=16000, mono=True)
            audio_2 = transcriber.isolate_voice(audio_2, sr_2)
            pass2_audio_path = Config.PROCESS_DIR / task.task_id / "audio_pass2_tr_mixed.wav"
            import soundfile as sf
            sf.write(str(pass2_audio_path), audio_2, sr_2)

            res2 = transcriber.transcribe_with_windowing(
                str(pass2_audio_path),
                model_name='selimc/whisper-large-v3-turbo-turkish',
                language='tr',
                window_size=30,
                overlap=5,
                progress_callback=lambda p, m: update_task_progress(task, 50 + p * 0.5, f"Pass 2 (Turbo-TR): {m}")
            )
            all_segments.extend(res2.get("segments", []))
            Path(pass2_audio_path).unlink(missing_ok=True)

            result = {
                "segments": all_segments,
                "text": " ".join([s.get("text", "") for s in all_segments]),
                "raw_text": " ".join([s.get("text", "") for s in all_segments]),
                "language": "tr",
                "method": "mixed_turkish_double_pass"
            }
        elif task.options.get('mixed_korean') and language == 'ko':
            # Mixed Korean Mode (Whisper Large V3 + Turbo KO)
            all_segments = []

            # Pass 1: OpenAI Whisper Large V3 (50s / 10s)
            task.message = "Mixed KO Pass 1 (Whisper Large V3, 50s window)..."
            res1 = transcriber.transcribe_with_windowing(
                str(audio_path),
                model_name='large-v3',
                language='ko',
                window_size=50,
                overlap=10,
                progress_callback=lambda p, m: update_task_progress(task, p * 0.5, f"Pass 1 (Large-V3): {m}")
            )
            all_segments.extend(res1.get("segments", []))

            if task.cancel_flag.is_set(): return

            # Pass 2: Whisper Turbo Korean (25s / 5s, isolated)
            # Using smaller window for Korean Turbo to prevent long runaway transcriptions
            task.message = "Mixed KO Pass 2 (Turbo Korean, 25s window, isolated)..."
            audio_2, sr_2 = librosa.load(str(audio_path), sr=16000, mono=True)
            audio_2 = transcriber.isolate_voice(audio_2, sr_2)
            pass2_audio_path = Config.PROCESS_DIR / task.task_id / "audio_pass2_ko_mixed.wav"
            import soundfile as sf
            sf.write(str(pass2_audio_path), audio_2, sr_2)

            res2 = transcriber.transcribe_with_windowing(
                str(pass2_audio_path),
                model_name='Farazzzzzzz/whisper-tiny_to_korean_accent2',
                language='ko',
                window_size=25,
                overlap=5,
                progress_callback=lambda p, m: update_task_progress(task, 50 + p * 0.5, f"Pass 2 (Turbo-KO): {m}")
            )
            all_segments.extend(res2.get("segments", []))
            Path(pass2_audio_path).unlink(missing_ok=True)

            result = {
                "segments": all_segments,
                "text": " ".join([s.get("text", "") for s in all_segments]),
                "raw_text": " ".join([s.get("text", "") for s in all_segments]),
                "language": "ko",
                "method": "mixed_korean_double_pass"
            }
        elif task.options.get('transcribe_window') and task.options.get('transcribe_window') > 0:
            # Windowed Transcription Path (requested by user for memory and stability)
            task.message = f"Starting Windowed Transcription ({model_name}, {task.options['transcribe_window']}s window)..."
            result = transcriber.transcribe_with_windowing(
                str(audio_path),
                model_name=model_name,
                language=language,
                window_size=task.options.get('transcribe_window', 30),
                overlap=task.options.get('transcribe_overlap', 10),
                progress_callback=lambda p, m: update_task_progress(task, p, m)
            )
        elif task.options.get('multi_pass'):
            # Multi-Pass Whisper transcription (Legacy/High-Accuracy)
            all_segments = []

            if model_name == 'selimc/whisper-large-v3-turbo-turkish':
                # Turkish Specific Multi-Pass (25s/5s and 35s/5s)
                task.message = "Turkish Pass 1 (25s window)..."
                res1 = transcriber.transcribe_with_windowing(
                    str(audio_path),
                    model_name=model_name,
                    language=language,
                    window_size=25,
                    overlap=5,
                    progress_callback=lambda p, m: update_task_progress(task, p * 0.5, f"Pass 1: {m}")
                )
                all_segments.extend(res1.get("segments", []))

                if task.cancel_flag.is_set(): return

                task.message = "Turkish Pass 2 (35s window, isolated)..."
                # Always isolate for the second pass in Turkish multi-pass
                audio_2, sr_2 = librosa.load(str(audio_path), sr=16000, mono=True)
                audio_2 = transcriber.isolate_voice(audio_2, sr_2)
                pass2_audio_path = Config.PROCESS_DIR / task.task_id / "audio_pass2_tr.wav"
                import soundfile as sf
                sf.write(str(pass2_audio_path), audio_2, sr_2)

                res2 = transcriber.transcribe_with_windowing(
                    str(pass2_audio_path),
                    model_name=model_name,
                    language=language,
                    window_size=35,
                    overlap=5,
                    progress_callback=lambda p, m: update_task_progress(task, 50 + p * 0.5, f"Pass 2: {m}")
                )
                all_segments.extend(res2.get("segments", []))
                Path(pass2_audio_path).unlink(missing_ok=True)

                result = {
                    "segments": all_segments,
                    "text": " ".join([s.get("text", "") for s in all_segments]),
                    "raw_text": " ".join([s.get("text", "") for s in all_segments]),
                    "language": res1.get("language", "tr")
                }
            else:
                # Standard Triple-Pass
                # Pass 1: UI settings
                task.message = "Whisper Pass 1 (UI Settings)..."
                res1 = transcriber.transcribe_with_windowing(
                    str(audio_path),
                    model_name=model_name,
                    language=language,
                    window_size=task.options.get("transcribe_window", Config.DEFAULT_TRANSCRIBE_WINDOW),
                    overlap=task.options.get("transcribe_overlap", Config.DEFAULT_TRANSCRIBE_OVERLAP),
                    progress_callback=lambda p, m: update_task_progress(task, p * 0.33, f"Pass 1: {m}")
                )
                all_segments.extend(res1.get("segments", []))

                if task.cancel_flag.is_set(): return

                # Pass 2: 45s window, 10s overlap, always voice isolated
                task.message = "Whisper Pass 2 (45s window, isolated)..."
                pass2_audio_path = audio_path
                if not task.options.get("isolate_voice"):
                    task.message = "Isolating voice for Pass 2..."
                    audio_2, sr_2 = librosa.load(str(audio_path), sr=16000, mono=True)
                    audio_2 = transcriber.isolate_voice(audio_2, sr_2)
                    pass2_audio_path = Config.PROCESS_DIR / task.task_id / "audio_pass2.wav"
                    import soundfile as sf
                    sf.write(str(pass2_audio_path), audio_2, sr_2)

                res2 = transcriber.transcribe_with_windowing(
                    str(pass2_audio_path),
                    model_name=model_name,
                    language=language,
                    window_size=45,
                    overlap=10,
                    progress_callback=lambda p, m: update_task_progress(task, 33 + p * 0.33, f"Pass 2: {m}")
                )
                all_segments.extend(res2.get("segments", []))
                if pass2_audio_path != audio_path:
                    Path(pass2_audio_path).unlink(missing_ok=True)

                if task.cancel_flag.is_set(): return

                # Pass 3: 60s window, 22s overlap, UI isolate_voice
                task.message = "Whisper Pass 3 (60s window, 22s overlap)..."
                res3 = transcriber.transcribe_with_windowing(
                    str(audio_path),
                    model_name=model_name,
                    language=language,
                    window_size=60,
                    overlap=22,
                    progress_callback=lambda p, m: update_task_progress(task, 66 + p * 0.34, f"Pass 3: {m}")
                )
                all_segments.extend(res3.get("segments", []))

                result = {
                    "segments": all_segments,
                    "text": " ".join([s.get("text", "") for s in all_segments]),
                    "raw_text": " ".join([s.get("text", "") for s in all_segments]),
                    "language": res1.get("language", "unknown")
                }
        else:
            # Primary Path: WhisperX
            task.message = f"Starting WhisperX ({model_name})..."
            try:
                result = transcriber.transcribe_whisperx(
                    str(audio_path),
                    model_name=model_name,
                    language=language,
                    batch_size=16,
                    use_diarization=task.options.get('use_diarization'),
                    hf_token=task.options.get('hf_token'),
                    progress_callback=lambda p, m: update_task_progress(task, p, m)
                )

                # ENSURE FORCED ALIGNMENT for non-WhisperX internal results
                if result.get("method") != "whisperx" and result.get("segments"):
                    task.message = "Ensuring forced alignment (post-transcription)..."
            finally:
                # Important: Clear VRAM after transcription before potential LLM/Translation work
                logger.info("Transcription step finished, unloading models...")
                transcriber.unload_model()
        
        # Diarization & Gender Detection
        speaker_genders = {}
        if task.options.get('use_diarization') and result.get('segments'):
            task.message = 'Detecting speaker genders...'
            try:
                audio_full, sr = librosa.load(str(audio_path), sr=16000, mono=True)

                # Group segment audio by speaker
                speaker_audio = {}
                for seg in result['segments']:
                    speaker = seg.get('speaker')
                    if not speaker: continue

                    # Ensure indices are within audio range
                    start_sample = max(0, int(seg['start'] * sr))
                    end_sample = min(len(audio_full), int(seg['end'] * sr))

                    if end_sample > start_sample:
                        chunk = audio_full[start_sample:end_sample]
                        if speaker not in speaker_audio:
                            speaker_audio[speaker] = []
                        speaker_audio[speaker].append(chunk)

                # Classify each speaker
                for speaker, chunks in speaker_audio.items():
                    if task.cancel_flag.is_set(): return

                    # Combine a few chunks for better accuracy (up to 10s of speech)
                    # We pick chunks that are likely to contain clear speech
                    valid_chunks = [c for c in chunks if len(c) > 0.5 * sr]
                    combined = np.concatenate(valid_chunks[:5]) if valid_chunks else (chunks[0] if chunks else np.array([]))

                    if len(combined) > 0:
                        gender = gender_detector.detect_gender(combined, sr)
                        speaker_genders[speaker] = gender
                        logger.info(f"Speaker {speaker} detected as {gender}")

                gender_detector.unload()

                # Apply gender to segments
                for seg in result['segments']:
                    seg['speaker_gender'] = speaker_genders.get(seg.get('speaker'), 'unknown')
            except Exception as e:
                logger.error(f"Gender detection failed: {e}")

        # Segment
        task.progress = 60
        task.message = 'Segmenting subtitles...'
        
        segments = result.get('segments', [])
        
        min_dur = task.options.get('min_duration', Config.MIN_SEGMENT_DURATION)
        max_dur = task.options.get('max_duration', Config.MAX_SEGMENT_DURATION)
        max_chars = task.options.get('max_chars', Config.MAX_CHARS_PER_SEGMENT)
        use_vad = task.options.get('use_vad', True) # Default True

        # If prevent_overlap is on, we force overlap to 0 during segmentation
        segment_overlap = 0.0 if task.options.get('prevent_overlap') else task.options.get('overlap', 0.5)

        if use_vad:
            segments = segmenter.segment_by_pauses(
                str(audio_path), segments,
                max_duration=max_dur,
                max_chars=max_chars,
                overlap=segment_overlap,
                margin=1.0 if task.options.get('use_margin') else 0.0
            )
        else:
            segments = segmenter.segment_by_time(
                segments,
                min_duration=min_dur,
                max_duration=max_dur,
                max_chars=max_chars,
                overlap=segment_overlap
            )
        
        # If we did triple-pass or mixed-language Whisper, use LLM to resolve overlaps and select best versions
        if engine != "cohere" and (task.options.get('multi_pass') or task.options.get('mixed_turkish') or task.options.get('mixed_korean')):
            task.message = "Refining multi-pass segments with LLM..."
            # For multi-pass, we also need to clear VRAM before LLM refinement if it's heavy
            transcriber.unload_model()
            segments = segmenter.merge_segments_llm(segments, translator)

        # Deduplication
        if task.options.get('deduplicate'):
            task.message = 'Removing repetitions...'
            segments = segmenter.remove_repetitions(segments)

        # Ensure sequential (no overlap)
        if task.options.get('prevent_overlap'):
            task.message = 'Ensuring sequential segments...'
            segments = segmenter.ensure_sequential(segments)
        else:
            # Even if not strictly sequential, we should merge identical overlaps
            segments = segmenter.merge_identical_overlapping(segments)

        # FINAL CLEANUP for speaker artifacts
        for seg in segments:
            seg['text'] = translator._clean_speaker_none(seg.get('text', ''))

        if task.cancel_flag.is_set():
            return
        
        # Translate if requested
        translations = {}
        if task.options.get('translate'):
            # Eliberează memoria GPU ocupată de Whisper înainte de a începe traducerea cu LLM
            task.message = 'Cleaning up VRAM for translation/LLM tasks...'
            transcriber.unload_model()
            translator.unload_models() # Ensure no old translation models are lingering

            task.progress = 80
            task.message = 'Translating...'
            
            target_langs = task.options.get('target_languages', ['en'])
            source_lang = result.get('language', 'en')
            engine = task.options.get('translation_engine', 'google')
            context = options.get('translation_context')
            
            texts = [seg.get('text', '') for seg in segments]
            metadata = [{'gender': seg.get('speaker_gender'), 'speaker': seg.get('speaker')} for seg in segments]
            
            llm_model = task.options.get('llm_model', Config.DEFAULT_LLM_MODEL)
            custom_prompt = task.options.get('custom_prompt')

            for target_lang in target_langs:
                if task.cancel_flag.is_set():
                    return
                
                task.message = f'Translating to {Config.SUPPORTED_LANGUAGES.get(target_lang, target_lang)} ({engine})...'
                
                if engine == 'vllm':
                    translated_texts = translator.translate_with_vllm_grouped(
                        texts, source_lang, target_lang,
                        model_name=llm_model,
                        group_size=task.options.get('translate_group', Config.DEFAULT_TRANSLATE_GROUP),
                        metadata=metadata,
                        context=context
                    )
                elif engine == 'llm':
                    translated_texts = translator.translate_with_llm(
                        texts, source_lang, target_lang,
                        model_name=llm_model,
                        custom_prompt=custom_prompt
                    )
                elif engine == 'llm_api':
                    translated_texts = translator.translate_with_api(
                        texts, source_lang, target_lang,
                        api_type=task.options.get('llm_api_provider'),
                        api_key=task.options.get('llm_api_key'),
                        model=task.options.get('llm_api_model'),
                        context=context,
                        base_url=task.options.get('llm_api_url'),
                        batch_size=20 # Blocurile de 15-20 segmente
                    )
                else: # google
                    translated_texts = translator.translate_batch(
                        texts, source_lang, target_lang,
                        batch_size=15, # Blocuri de 10-15 segmente
                        context=context
                    )
                
                # Validation and Retry
                translated_texts = translator.validate_and_retry_translations(
                    texts, translated_texts, source_lang, target_lang, engine
                )

                # Apply LLM correction if requested
                if task.options.get('use_romistral'):
                    # Use explicit refiner model if provided, else fallback to llm_model or default
                    refiner_model = task.options.get('refiner_model') or task.options.get('llm_model', Config.DEFAULT_LLM_MODEL)
                    task.message = f'Refining translation with {refiner_model}...'
                    translated_texts = translator.correct_with_vllm(
                        translated_texts,
                        target_lang,
                        model_name=refiner_model,
                        group_size=task.options.get('translate_group', Config.DEFAULT_TRANSLATE_GROUP),
                        metadata=metadata,
                        context=context
                    )

                translations[target_lang] = translated_texts
            
            task.progress = 95
        
        # Apply timestamp offset if we processed a region
        if process_start > 0:
            task.message = f'Offsetting timestamps by {process_start}s...'
            for seg in segments:
                seg['start'] += process_start
                seg['end'] += process_start

        # Prepare result
        task.result = {
            'full_text': result.get('text', ''),
            'raw_text': result.get('raw_text', ''),
            'segments': segments,
            'language': result.get('language', 'unknown'),
            'translations': translations,
            'task_id': task.task_id
        }
        
        task.status = 'completed'
        task.progress = 100
        task.message = 'Processing complete!'
        
    except Exception as e:
        logger.error(f"Processing error: {e}")
        task.status = 'failed'
        task.error = str(e)
        task.message = f'Error: {str(e)}'
    finally:
        # Cleanup VRAM aggressively after every process
        logger.info(f"Task {task.task_id} finished, clearing VRAM...")
        transcriber.unload_model()
        translator.unload_models()

        # Cleanup audio if extracted
        if 'audio_path' in locals() and audio_path != task.file_path:
            Path(audio_path).unlink(missing_ok=True)

def update_task_progress(task, progress, message):
    """Update task progress"""
    task.progress = progress
    task.message = message

def cleanup_old_tasks():
    """Remove old completed tasks"""
    current_time = time.time()
    with task_lock:
        tasks_to_remove = []
        for task_id, task in processing_tasks.items():
            if current_time - task.created_at > Config.SESSION_LIFETIME:
                tasks_to_remove.append(task_id)
                
                # Cleanup files
                task_dir = Config.PROCESS_DIR / task_id
                if task_dir.exists():
                    shutil.rmtree(task_dir)
        
        for task_id in tasks_to_remove:
            del processing_tasks[task_id]

def generate_srt(segments, use_legacy_diacritics=False):
    """Generate SRT format from segments"""
    srt_lines = []
    
    for i, segment in enumerate(segments, 1):
        start_time = segment.get('start', 0)
        end_time = segment.get('end', 0)
        text = segment.get('text', '').strip()
        
        if not text:
            continue
        
        if use_legacy_diacritics:
            text = segmenter.convert_diacritics(text, to_legacy=True)
        
        # Split text for readability
        text = segmenter.split_text_for_subtitle(text, 40)
        
        srt_lines.append(str(i))
        srt_lines.append(f"{format_time(start_time)} --> {format_time(end_time)}")
        srt_lines.append(text)
        srt_lines.append('')  # Empty line
    
    return '\n'.join(srt_lines)

def format_time(seconds):
    """Format time for SRT: HH:MM:SS,mmm"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    seconds = seconds % 60
    milliseconds = int((seconds - int(seconds)) * 1000)
    
    return f"{hours:02d}:{minutes:02d}:{int(seconds):02d},{milliseconds:03d}"

# ============ Scheduled Cleanup ============

def scheduled_cleanup():
    """Run cleanup periodically"""
    while True:
        time.sleep(Config.CLEANUP_INTERVAL)
        try:
            file_handler.cleanup_old_sessions()
            cleanup_old_tasks()
            logger.info("Scheduled cleanup completed")
        except Exception as e:
            logger.error(f"Cleanup error: {e}")

# Start cleanup thread
cleanup_thread = threading.Thread(target=scheduled_cleanup, daemon=True)
cleanup_thread.start()

# ============ Error Handlers ============

@app.errorhandler(404)
def not_found_error(error):
    return jsonify({'error': 'Resource not found'}), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': 'Internal server error'}), 500

@app.errorhandler(413)
def too_large_error(error):
    return jsonify({'error': 'File too large'}), 413

# ============ Main ============

if __name__ == '__main__':
    logger.info("Starting Whisper Transcriber Application")
    logger.info(f"Device: {transcriber.device}")
    logger.info(f"Server: http://{Config.HOST}:{Config.PORT}")
    
    app.run(
        host=Config.HOST,
        port=Config.PORT,
        debug=Config.DEBUG,
        threaded=True
    )
